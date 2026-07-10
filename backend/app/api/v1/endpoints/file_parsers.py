"""
File Parsing Utilities.

Functions to parse and extract data from various file formats (CSV, Excel, Numbers, TXT, PDF).
"""
import io
import docx
import tempfile
import os
from openpyxl import load_workbook
from pypdf import PdfReader
from numbers_parser import Document

def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """
    Extracts plain text from document files (PDF, DOCX, TXT).
    
    Args:
        filename (str): The name of the file being uploaded.
        content (bytes): The raw bytes of the file.
        
    Returns:
        str: The extracted plain text. Returns empty string if parsing fails.
    """
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    try:
        if ext in ['txt', 'csv']:
            return content.decode("utf-8")
        
        elif ext in ['docx', 'doc']:
            doc = docx.Document(io.BytesIO(content))
            return "\n".join([para.text for para in doc.paragraphs])
            
        elif ext in ['xlsx', 'xls']:
            wb = load_workbook(filename=io.BytesIO(content), data_only=True)
            text = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) for cell in row if cell is not None])
                    if row_text:
                        text.append(row_text)
            return "\n".join(text)
            
        elif ext == 'pdf':
            reader = PdfReader(io.BytesIO(content))
            text = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text.append(extracted)
            return "\n".join(text)
            
        elif ext == 'numbers':
            with tempfile.NamedTemporaryFile(delete=False, suffix=".numbers") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
                
            try:
                doc = Document(tmp_path)
                text = []
                for sheet in doc.sheets:
                    for table in sheet.tables:
                        for row in table.rows():
                            row_text = " ".join([str(cell.value) for cell in row if cell.value is not None])
                            if row_text:
                                text.append(row_text)
                return "\n".join(text)
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            
        elif ext == 'pages':
            return "Pages file detected. Native parsing is not supported. Please convert to PDF or Word for better results."
            
        else:
            return content.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"Error parsing {filename}: {e}")
        return f"Error extracting text from {filename}"

def extract_tabular_data_from_bytes(filename: str, content: bytes) -> list:
    """
    Parses tabular data files (CSV, Excel, Numbers) into a list of dictionaries.
    
    Args:
        filename (str): The name of the tabular file.
        content (bytes): The raw bytes of the file.
        
    Returns:
        list[dict]: A list of records, where each record is a dictionary mapping 
                    column headers to row values.
    """
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    rows = []
    
    try:
        if ext in ['xlsx', 'xls']:
            wb = load_workbook(filename=io.BytesIO(content), data_only=True)
            sheet = wb.active
            headers = [str(cell.value).strip() if cell.value else "" for cell in sheet[1]]
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if any(row):
                    row_dict = {headers[i]: str(row[i]) if row[i] is not None else "" for i in range(len(headers))}
                    rows.append(row_dict)
                    
        elif ext == 'numbers':
            with tempfile.NamedTemporaryFile(delete=False, suffix=".numbers") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
                
            try:
                doc = Document(tmp_path)
                sheet = doc.sheets[0]
                table = sheet.tables[0]
                header_row = table.rows()[0]
                headers = [str(cell.value).strip() if cell.value else "" for cell in header_row]
                
                for row in table.rows()[1:]:
                    if any(cell.value is not None for cell in row):
                        row_dict = {headers[i]: str(row[i].value) if i < len(row) and row[i].value is not None else "" for i in range(len(headers))}
                        rows.append(row_dict)
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                    
        elif ext == 'csv':
            import csv
            reader = csv.DictReader(io.StringIO(content.decode("utf-8")))
            for row in reader:
                rows.append(dict(row))
                
        return rows
    except Exception as e:
        print(f"Error parsing tabular data {filename}: {e}")
        return []
